"""PDF / DOCX / TXT text extraction."""

from __future__ import annotations

import io

from app.core.config import get_settings
from app.core.errors import DocumentParseError

SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".txt", ".md")


def extract_text(filename: str, data: bytes) -> str:
    settings = get_settings()
    if not data:
        raise DocumentParseError("The uploaded file is empty.")
    if len(data) > settings.max_upload_bytes:
        raise DocumentParseError(
            f"File is too large (limit {settings.max_upload_bytes // (1024 * 1024)} MB)."
        )

    lower = filename.lower()
    if lower.endswith(".pdf"):
        text = _extract_pdf(data)
    elif lower.endswith(".docx"):
        text = _extract_docx(data)
    elif lower.endswith((".txt", ".md")):
        text = _extract_txt(data)
    elif lower.endswith(".doc"):
        raise DocumentParseError(
            "Legacy .doc files are not supported. Please upload PDF, DOCX or TXT."
        )
    else:
        raise DocumentParseError(
            "Unsupported file type. Please upload a PDF, DOCX or TXT file."
        )

    text = _normalize(text)
    if len(text) < 20:
        raise DocumentParseError(
            "No readable text was found in this file. If it is a scanned PDF, "
            "please upload a text-based version."
        )
    return text[: settings.max_document_chars]


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        return "\n\n".join((page.extract_text() or "") for page in reader.pages)
    except DocumentParseError:
        raise
    except Exception as exc:  # noqa: BLE001 - surfaced to the user as a clean error
        raise DocumentParseError(f"Could not read this PDF: {exc}") from exc


def _extract_docx(data: bytes) -> str:
    try:
        import docx

        document = docx.Document(io.BytesIO(data))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(parts)
    except Exception as exc:  # noqa: BLE001
        raise DocumentParseError(f"Could not read this DOCX file: {exc}") from exc


def _extract_txt(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentParseError("Could not decode this text file.")


def _normalize(text: str) -> str:
    lines = [line.rstrip() for line in text.splitlines()]
    out: list[str] = []
    blank = 0
    for line in lines:
        if not line.strip():
            blank += 1
            if blank > 1:
                continue
        else:
            blank = 0
        out.append(line)
    return "\n".join(out).strip()
